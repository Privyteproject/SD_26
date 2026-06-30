"""Workflow de génération documentaire en 2 temps (preview -> submit).

- Jeton de prévisualisation signé HMAC-SHA256 (anti-falsification).
- Données de preview stockées dans Redis (TTL court), jamais en base à cette étape.
- Rendu via templates Jinja2 `app/templates/documents/{type}.html.j2` (HTML imprimable).
"""

import hashlib
import hmac
import os
import re
import secrets
from datetime import date, datetime

from app.core.config import settings
from app.services.document_types import DOCUMENT_TYPES, label_of

_TEMPLATES_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "templates")
_ENV = None


def _format_date(value):
    if isinstance(value, (date, datetime)):
        return value.strftime("%d/%m/%Y")
    return str(value) if value is not None else ""


def _anciennete_annees(value):
    """Ancienneté en années (entier) depuis une date d'entrée. 0 si inconnue/non datée."""
    if isinstance(value, str):
        try:
            value = datetime.fromisoformat(value).date()
        except ValueError:
            return 0
    if isinstance(value, datetime):
        value = value.date()
    if not isinstance(value, date):
        return 0
    today = date.today()
    return max(0, today.year - value.year - ((today.month, today.day) < (value.month, value.day)))


def _env():
    global _ENV
    if _ENV is not None:
        return _ENV
    from jinja2 import Environment, FileSystemLoader, select_autoescape

    _ENV = Environment(
        # On cherche à la racine ET dans documents/ pour que `{% extends "_base.html.j2" %}`
        # résolve le base situé dans documents/.
        loader=FileSystemLoader([_TEMPLATES_DIR, os.path.join(_TEMPLATES_DIR, "documents")]),
        autoescape=select_autoescape(["html", "j2"]),
    )
    _ENV.filters["format_date"] = _format_date
    _ENV.filters["anciennete_annees"] = _anciennete_annees
    return _ENV


# ───────────── Jeton signé ─────────────
def make_token() -> tuple[str, str]:
    nonce = secrets.token_urlsafe(16)
    sig = hmac.new(settings.DOC_PREVIEW_SECRET.encode(), nonce.encode(), hashlib.sha256).hexdigest()[:24]
    return f"{nonce}.{sig}", nonce


def verify_token(token: str) -> str | None:
    try:
        nonce, sig = token.split(".", 1)
    except (ValueError, AttributeError):
        return None
    expected = hmac.new(settings.DOC_PREVIEW_SECRET.encode(), nonce.encode(), hashlib.sha256).hexdigest()[:24]
    return nonce if hmac.compare_digest(sig, expected) else None


# ───────────── Rendu ─────────────
def _html_to_text(html: str) -> str:
    import html as _h
    # Retire d'abord les blocs <style>/<script> (leur contenu n'est pas du texte).
    txt = re.sub(r"(?is)<(style|script)[^>]*>.*?</\1>", "", html)
    txt = re.sub(r"(?i)</p>|<br\s*/?>|</h1>|</div>", "\n", txt)
    txt = re.sub(r"<[^>]+>", "", txt)
    txt = _h.unescape(txt)
    txt = re.sub(r"\n{3,}", "\n\n", txt)
    return "\n".join(line.strip() for line in txt.splitlines()).strip()


def _fallback_html(label: str, emp: dict, additional: dict) -> str:
    nom = f"{emp.get('prenom', '')} {emp.get('nom', '')}".strip()
    objet = additional.get("objet") or additional.get("motif") or ""
    return (
        f"<div style='font-family:Georgia,serif'><h2 style='text-align:center'>{label}</h2>"
        f"<p>{settings.COMPANY_NAME} certifie les informations relatives à "
        f"<b>{nom}</b> ({emp.get('poste') or '—'}).</p>"
        f"{('<p>Objet : ' + objet + '</p>') if objet else ''}"
        f"<p style='color:#888;font-size:12px'>{_format_date(date.today())} — "
        f"Document généré par la plateforme RH</p></div>"
    )


def render(doc_type: str, emp: dict, additional: dict | None = None) -> tuple[str, str, str]:
    """Renvoie (document_name, html_preview, text_content)."""
    additional = additional or {}
    label = label_of(doc_type)
    document_name = f"{doc_type}_{emp.get('matricule', 'doc')}.txt"
    ctx = {
        "employee": emp,
        "additional": additional,
        "company": {"nom": settings.COMPANY_NAME, "adresse": settings.COMPANY_ADDRESS},
        "date_generation": _format_date(date.today()),
        "label": label,
        "requires_rh_validation": DOCUMENT_TYPES.get(doc_type, {}).get("requires_rh_validation", False),
        **additional,  # champs spécifiques accessibles directement ({{ objet }}, {{ motif }}…)
    }
    try:
        html = _env().get_template(f"documents/{doc_type}.html.j2").render(**ctx)
    except Exception:
        html = _fallback_html(label, emp, additional)
    return document_name, html, _html_to_text(html)
import io
import re
from pypdf import PdfReader, PdfWriter
from docx import Document
from jinja2 import Template

def resolve_context_value(ctx, key):
    # key e.g., "employee.nom", "employee_nom", "nom_complet"
    key = key.strip().lower()
    
    # Handle specific computed fields like nom_complet
    if "nom_complet" in key:
        emp = ctx.get("employee", {})
        return f"{emp.get('prenom', '')} {emp.get('nom', '')}".strip()
        
    parts = key.replace('_', '.').split('.')
    val = ctx
    for p in parts:
        if isinstance(val, dict):
            # Case insensitive key search
            found = False
            for k, v in val.items():
                if k.lower() == p:
                    val = v
                    found = True
                    break
            if not found:
                return ""
        else:
            return ""
    return str(val) if val is not None else ""

def fill_docx_template(binary_content, ctx):
    doc = Document(io.BytesIO(binary_content))
    
    def process_runs(runs):
        text = "".join(run.text for run in runs)
        if "{{" in text and "}}" in text:
            template = Template(text)
            rendered = template.render(**ctx)
            if runs:
                runs[0].text = rendered
                for run in runs[1:]:
                    run.text = ""
                    
    for paragraph in doc.paragraphs:
        process_runs(paragraph.runs)
        
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    process_runs(paragraph.runs)
                    
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()

def fill_pdf_template(binary_content, ctx):
    reader = PdfReader(io.BytesIO(binary_content))
    writer = PdfWriter()
    
    # Get form fields
    fields = reader.get_fields()
    field_updates = {}
    if fields:
        for field_name in fields.keys():
            val = resolve_context_value(ctx, field_name)
            if val:
                field_updates[field_name] = val
                
    for page in reader.pages:
        writer.add_page(page)
        
    if field_updates:
        writer.update_page_form_field_values(
            writer.pages[0], field_updates, auto_regenerate=False
        )
        
    # Flatten PDF to prevent further editing
    for page in writer.pages:
        if "/Annots" in page:
            for annot in page["/Annots"]:
                annot_obj = annot.get_object()
                if annot_obj.get("/Subtype") == "/Widget":
                    annot_obj.update({
                        # Set readonly flag (bit 1)
                        "/Ff": 1
                    })

    out = io.BytesIO()
    writer.write(out)
    return out.getvalue()

def docx_to_html_preview(binary_content, document_name):
    doc = Document(io.BytesIO(binary_content))
    html = [
        f"<div style='font-family:Arial,sans-serif;max-width:800px;margin:0 auto;padding:20px;background:white;box-shadow:0 0 10px rgba(0,0,0,0.1);'>",
        f"<h3 style='text-align:center;color:#333;margin-bottom:20px;'>Aperçu du document : {document_name}</h3>"
    ]
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            html.append(f"<p style='margin-bottom:10px;line-height:1.5;'>{text}</p>")
            
    for table in doc.tables:
        html.append("<table style='width:100%;border-collapse:collapse;margin-bottom:15px;'>")
        for row in table.rows:
            html.append("<tr>")
            for cell in row.cells:
                html.append(f"<td style='border:1px solid #ddd;padding:8px;'>{cell.text.strip()}</td>")
            html.append("</tr>")
        html.append("</table>")
        
    html.append("</div>")
    return "\n".join(html)

